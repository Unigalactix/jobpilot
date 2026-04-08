"""
resume_reader.py  —  reads .txt, .docx, and .pdf resume files
"""

import os
from pathlib import Path


RESUMES_DIR = Path(__file__).parent / "resumes"


def get_resume_list() -> list[dict]:
    """Return list of resume files with metadata."""
    files = []
    if not RESUMES_DIR.exists():
        return files
    for f in sorted(RESUMES_DIR.iterdir()):
        if f.suffix.lower() in (".txt", ".docx", ".pdf"):
            files.append({
                "name": f.name,
                "size": f.stat().st_size,
                "ext":  f.suffix.lower(),
            })
    return files


def read_resume(filename: str) -> str:
    """Read and return resume text from file."""
    path = RESUMES_DIR / filename
    if not path.exists():
        return ""

    ext = path.suffix.lower()

    if ext == ".txt":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(str(path))
            lines = []
            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    lines.append(t)
            return "\n".join(lines)
        except Exception as e:
            print(f"[reader] docx error: {e}")
            return ""

    if ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            text = "\n".join(p.extract_text() or "" for p in reader.pages)
            if text.strip():
                return text
        except Exception:
            pass
        try:
            import pdfplumber
            lines = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        lines.append(t)
            return "\n".join(lines)
        except Exception as e:
            print(f"[reader] pdf error: {e}")
            return ""

    return ""


def save_tailored_resume(filename: str, content: str) -> str:
    """Save tailored resume as .txt in generated/ folder."""
    out_dir = Path(__file__).parent / "generated"
    out_dir.mkdir(exist_ok=True)
    base = Path(filename).stem
    out_path = out_dir / f"{base}_tailored.txt"
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(content)
    return str(out_path)


def save_tailored_pdf(filename: str, content: str, max_pages: int = 0) -> str:
    """
    Save tailored resume as PDF matching the original format exactly.
    max_pages=0 → no limit; max_pages=1 or 2 → auto-shrink to fit.
    """
    import re
    import io

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
        )
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

        out_dir = Path(__file__).parent / "generated"
        out_dir.mkdir(exist_ok=True)
        base = Path(filename).stem
        out_path = out_dir / f"{base}_tailored.pdf"

        # ── Exact measurements from original PDF ──────────────────────────────
        L_MARGIN   = 28.8   # pt
        R_MARGIN   = 29.5   # pt
        T_MARGIN   = 33.7   # pt
        B_MARGIN   = 33.8   # pt
        PAGE_W     = 612    # pt (letter)
        TEXT_W     = PAGE_W - L_MARGIN - R_MARGIN   # 553.7 pt

        doc = SimpleDocTemplate(
            str(out_path),
            pagesize=letter,
            leftMargin=L_MARGIN,
            rightMargin=R_MARGIN,
            topMargin=T_MARGIN,
            bottomMargin=B_MARGIN,
        )

        # ── Styles ────────────────────────────────────────────────────────────
        BLACK = colors.black

        name_style = ParagraphStyle("name",
            fontName="Times-Bold", fontSize=20, leading=24,
            alignment=TA_CENTER, spaceAfter=1, textColor=BLACK)

        contact_style = ParagraphStyle("contact",
            fontName="Times-Roman", fontSize=10, leading=13,
            alignment=TA_CENTER, spaceAfter=5, textColor=BLACK)

        section_style = ParagraphStyle("section",
            fontName="Times-Bold", fontSize=12, leading=14,
            alignment=TA_LEFT, spaceBefore=7, spaceAfter=3, textColor=BLACK)

        # Company | Role line (left) and Date (right) handled via Table
        company_left = ParagraphStyle("company_left",
            fontName="Times-Bold", fontSize=10, leading=12,
            alignment=TA_LEFT, textColor=BLACK)

        company_right = ParagraphStyle("company_right",
            fontName="Times-Bold", fontSize=10, leading=12,
            alignment=TA_RIGHT, textColor=BLACK)

        # Sub-location line (e.g. "Redmond, WA | Project: ...")
        subloc_style = ParagraphStyle("subloc",
            fontName="Times-Italic", fontSize=9, leading=11,
            alignment=TA_LEFT, spaceAfter=1, textColor=BLACK)

        body_style = ParagraphStyle("body",
            fontName="Times-Roman", fontSize=9, leading=12,
            alignment=TA_LEFT, spaceAfter=1, textColor=BLACK)

        # Bullets: • at 13.3pt, text continuation at 21.9pt from left margin
        bullet_style = ParagraphStyle("bullet",
            fontName="Times-Roman", fontSize=9, leading=12,
            alignment=TA_LEFT,
            leftIndent=21.9, firstLineIndent=-8.6,
            spaceAfter=1, textColor=BLACK)

        italic_style = ParagraphStyle("italic",
            fontName="Times-Italic", fontSize=9, leading=12,
            alignment=TA_LEFT, spaceAfter=1, textColor=BLACK)

        skills_label_style = ParagraphStyle("skills_label",
            fontName="Times-Bold", fontSize=9, leading=12,
            alignment=TA_LEFT, spaceAfter=1, textColor=BLACK)

        # ── Helpers ───────────────────────────────────────────────────────────
        DATE_RE = re.compile(
            r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s*\d{4}'
            r'\s*[–—\-]+\s*(?:Present|[A-Z][a-z]+\s*\d{0,4}))',
            re.IGNORECASE
        )

        def _is_company_line(line):
            """True if line looks like 'Company | Role   Date'."""
            return bool(DATE_RE.search(line)) and ("|" in line or len(line.split()) < 8)

        def _is_section_header(line):
            stripped = line.strip()
            if len(stripped) < 3:
                return False
            # ALL CAPS or common section names
            return stripped.replace(" ", "").replace("&", "").replace("/", "").isupper()

        def _split_company_date(line):
            m = DATE_RE.search(line)
            if m:
                left  = line[:m.start()].strip().rstrip("|,").strip()
                right = m.group(1).strip()
                return left, right
            return line.strip(), ""

        def _make_company_row(left_text, right_text):
            """Table row: bold left (company|role) + bold right (date)."""
            row = [[
                Paragraph(left_text.replace("&", "&amp;").replace("<", "&lt;"), company_left),
                Paragraph(right_text.replace("&", "&amp;"), company_right),
            ]]
            t = Table(row, colWidths=[TEXT_W * 0.72, TEXT_W * 0.28])
            t.setStyle(TableStyle([
                ("VALIGN",       (0, 0), (-1, -1), "BOTTOM"),
                ("LEFTPADDING",  (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING",   (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING",(0, 0), (-1, -1), 0),
            ]))
            return t

        def _escape(s):
            return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        # ── Parse and build story ─────────────────────────────────────────────
        raw_lines = content.split("\n")
        story     = []

        # Find first non-empty line index
        first_content = next((i for i, l in enumerate(raw_lines) if l.strip()), 0)
        header_done = False     # True after contact line is emitted
        in_education = False
        i = first_content

        while i < len(raw_lines):
            line = raw_lines[i].strip()
            i += 1

            if not line:
                continue

            # ── Name (first 1-2 non-empty lines before contact) ──────────────
            if not header_done and ("@" not in line) and ("|" not in line or i - first_content <= 2):
                # Could be name — check it's not a section header starting immediately
                if not _is_section_header(line) or i - first_content <= 2:
                    story.append(Paragraph(_escape(line), name_style))
                    continue

            # ── Contact line ─────────────────────────────────────────────────
            if not header_done and ("|" in line or "@" in line):
                story.append(Paragraph(_escape(line.replace("|", "  |  ")), contact_style))
                story.append(HRFlowable(
                    width=TEXT_W, thickness=0.8, color=BLACK,
                    spaceBefore=0, spaceAfter=5, hAlign="LEFT"
                ))
                header_done = True
                continue

            # ── Section headers ───────────────────────────────────────────────
            if _is_section_header(line):
                in_education = "EDUCATION" in line.upper()
                story.append(Paragraph(_escape(line), section_style))
                story.append(HRFlowable(
                    width=TEXT_W, thickness=0.5, color=colors.HexColor("#999999"),
                    spaceBefore=0, spaceAfter=3, hAlign="LEFT"
                ))
                continue

            # ── Company | Role line with date ─────────────────────────────────
            if _is_company_line(line):
                left, right = _split_company_date(line)
                story.append(_make_company_row(_escape(left), _escape(right)))
                continue

            # ── Bullet points ─────────────────────────────────────────────────
            if line.startswith(("•", "●", "▪", "*")) or (line.startswith("-") and len(line) > 2):
                text = line.lstrip("•●▪*- ").strip()
                story.append(Paragraph(f"• {_escape(text)}", bullet_style))
                continue

            # ── Education degree line (italic) ────────────────────────────────
            if in_education and any(kw in line for kw in
                    ["Master", "Bachelor", "Doctor", "PhD", "MS ", "BS ", "MBA",
                     "Science", "Engineering", "Arts", "Technology"]):
                story.append(Paragraph(_escape(line), italic_style))
                continue

            # ── Skills lines (bold label: value) ─────────────────────────────
            if ":" in line and not line.startswith("•"):
                colon = line.index(":")
                label = line[:colon + 1]
                value = line[colon + 1:]
                text  = f"<b>{_escape(label)}</b>{_escape(value)}"
                story.append(Paragraph(text, body_style))
                continue

            # ── Default body ──────────────────────────────────────────────────
            story.append(Paragraph(_escape(line), body_style))

        def _build_story_with_scale(scale):
            """Rebuild story with scaled font sizes for fit-to-page."""
            def sp(base): return ParagraphStyle
            scaled_name    = ParagraphStyle("name_s",    fontName="Times-Bold",   fontSize=20*scale,   leading=24*scale,  alignment=TA_CENTER, spaceAfter=1,        textColor=BLACK)
            scaled_contact = ParagraphStyle("contact_s", fontName="Times-Roman",  fontSize=10*scale,   leading=13*scale,  alignment=TA_CENTER, spaceAfter=5*scale,  textColor=BLACK)
            scaled_section = ParagraphStyle("section_s", fontName="Times-Bold",   fontSize=12*scale,   leading=14*scale,  alignment=TA_LEFT,   spaceBefore=7*scale, spaceAfter=3*scale, textColor=BLACK)
            scaled_co_l    = ParagraphStyle("co_l_s",    fontName="Times-Bold",   fontSize=10*scale,   leading=12*scale,  alignment=TA_LEFT,   textColor=BLACK)
            scaled_co_r    = ParagraphStyle("co_r_s",    fontName="Times-Bold",   fontSize=10*scale,   leading=12*scale,  alignment=TA_RIGHT,  textColor=BLACK)
            scaled_body    = ParagraphStyle("body_s",    fontName="Times-Roman",  fontSize=9*scale,    leading=12*scale,  alignment=TA_LEFT,   spaceAfter=1,        textColor=BLACK)
            scaled_bullet  = ParagraphStyle("bullet_s",  fontName="Times-Roman",  fontSize=9*scale,    leading=12*scale,  alignment=TA_LEFT,   leftIndent=21.9*scale, firstLineIndent=-8.6*scale, spaceAfter=1, textColor=BLACK)
            scaled_italic  = ParagraphStyle("italic_s",  fontName="Times-Italic", fontSize=9*scale,    leading=12*scale,  alignment=TA_LEFT,   spaceAfter=1,        textColor=BLACK)
            scaled_subloc  = ParagraphStyle("subloc_s",  fontName="Times-Italic", fontSize=9*scale,    leading=11*scale,  alignment=TA_LEFT,   spaceAfter=1,        textColor=BLACK)

            new_story = []
            _hd = False
            _ined = False
            raw = content.split("\n")
            fc = next((i for i, l in enumerate(raw) if l.strip()), 0)
            idx = fc
            while idx < len(raw):
                ln = raw[idx].strip(); idx += 1
                if not ln: continue
                if not _hd and "@" not in ln and ("|" not in ln or idx - fc <= 2):
                    if not (ln.replace(" ","").replace("&","").replace("/","").isupper()) or idx - fc <= 2:
                        new_story.append(Paragraph(_escape(ln), scaled_name)); continue
                if not _hd and ("|" in ln or "@" in ln):
                    new_story.append(Paragraph(_escape(ln.replace("|","  |  ")), scaled_contact))
                    new_story.append(HRFlowable(width=TEXT_W, thickness=0.8*scale, color=BLACK, spaceBefore=0, spaceAfter=5*scale, hAlign="LEFT"))
                    _hd = True; continue
                if ln.replace(" ","").replace("&","").replace("/","").isupper() and len(ln) > 2:
                    _ined = "EDUCATION" in ln.upper()
                    new_story.append(Paragraph(_escape(ln), scaled_section))
                    new_story.append(HRFlowable(width=TEXT_W, thickness=0.5*scale, color=colors.HexColor("#999999"), spaceBefore=0, spaceAfter=3*scale, hAlign="LEFT"))
                    continue
                if bool(DATE_RE.search(ln)) and ("|" in ln or len(ln.split()) < 8):
                    lft, rgt = _split_company_date(ln)
                    row = [[Paragraph(_escape(lft), scaled_co_l), Paragraph(_escape(rgt), scaled_co_r)]]
                    t = Table(row, colWidths=[TEXT_W*0.72, TEXT_W*0.28])
                    t.setStyle(TableStyle([("VALIGN",(0,0),(-1,-1),"BOTTOM"),("LEFTPADDING",(0,0),(-1,-1),0),("RIGHTPADDING",(0,0),(-1,-1),0),("TOPPADDING",(0,0),(-1,-1),3*scale),("BOTTOMPADDING",(0,0),(-1,-1),0)]))
                    new_story.append(t); continue
                if ln.startswith(("•","●","▪","*")) or (ln.startswith("-") and len(ln)>2):
                    new_story.append(Paragraph(f"• {_escape(ln.lstrip('•●▪*- ').strip())}", scaled_bullet)); continue
                if _ined and any(kw in ln for kw in ["Master","Bachelor","Doctor","PhD","MS ","BS ","MBA","Science","Engineering","Arts","Technology"]):
                    new_story.append(Paragraph(_escape(ln), scaled_italic)); continue
                if ":" in ln and not ln.startswith("•"):
                    colon = ln.index(":"); new_story.append(Paragraph(f"<b>{_escape(ln[:colon+1])}</b>{_escape(ln[colon+1:])}", scaled_body)); continue
                new_story.append(Paragraph(_escape(ln), scaled_body))
            return new_story

        # ── Build and auto-fit to max_pages ──────────────────────────────────
        def _count_pages(stry):
            buf = io.BytesIO()
            tmp = SimpleDocTemplate(buf, pagesize=letter,
                leftMargin=L_MARGIN, rightMargin=R_MARGIN,
                topMargin=T_MARGIN,  bottomMargin=B_MARGIN)
            tmp.build(stry)
            buf.seek(0)
            from pypdf import PdfReader
            return len(PdfReader(buf).pages)

        if max_pages > 0:
            for scale in [1.0, 0.95, 0.90, 0.85, 0.80, 0.75]:
                trial_story = _build_story_with_scale(scale)
                if _count_pages(trial_story) <= max_pages:
                    story = trial_story
                    break
            else:
                story = _build_story_with_scale(0.75)

        doc.build(story)
        return str(out_path)

    except Exception as e:
        print(f"[pdf export] Error: {e}")
        import traceback; traceback.print_exc()
        return save_tailored_resume(filename, content)


def save_tailored_docx(filename: str, content: str) -> str:
    """Save tailored resume as .docx in generated/ folder."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        out_dir = Path(__file__).parent / "generated"
        out_dir.mkdir(exist_ok=True)
        base = Path(filename).stem
        out_path = out_dir / f"{base}_tailored.docx"

        doc = Document()
        # Set margins
        for section in doc.sections:
            section.top_margin    = Pt(36)
            section.bottom_margin = Pt(36)
            section.left_margin   = Pt(54)
            section.right_margin  = Pt(54)

        for line in content.split("\n"):
            line = line.rstrip()
            if not line:
                doc.add_paragraph("")
                continue
            p = doc.add_paragraph()
            # Detect section headers (ALL CAPS lines)
            if line.isupper() and len(line) > 3:
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after  = Pt(2)
            elif line.startswith("•") or line.startswith("-"):
                run = p.add_run(line)
                run.font.size = Pt(9.5)
                p.paragraph_format.left_indent = Pt(14)
                p.paragraph_format.space_after = Pt(1)
            else:
                run = p.add_run(line)
                run.font.size = Pt(10)
                p.paragraph_format.space_after = Pt(2)

        doc.save(str(out_path))
        return str(out_path)
    except Exception as e:
        print(f"[docx export] Error: {e}")
        return save_tailored_resume(filename, content)
